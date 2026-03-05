from docx import Document
from docx.shared import RGBColor
from google import genai
from dotenv import load_dotenv
from bs4 import BeautifulSoup
import os
from docx_parser_converter import docx_to_html
import datetime
from docx2pdf import convert
import os
import subprocess


load_dotenv()

client = genai.Client(api_key=os.environ["GEMINI_KEY"])
chat = client.chats.create(model="gemini-2.5-flash-lite-preview-09-2025")

def convert_docx_to_pdf(input_file, output_dir=None):
    if output_dir is None:
        # Use the same directory as the input file
        output_dir = os.path.dirname(os.path.abspath(input_file))

    # Ensure the output directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # The command to run LibreOffice in headless mode
    command = [
        'libreoffice',
        '--headless',
        '--convert-to',
        'pdf',
        '--outdir',
        output_dir,
        input_file
    ]

    try:
        # Run the command
        subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e.stderr.decode()}")
    except FileNotFoundError:
        print("Error: libreoffice executable not found. Make sure LibreOffice is installed and in your PATH.")

def create_resume(filename, company_name=None, company_address=None, company_code=None, description=None):
    """
    Create a customized resume using Gemini
    """
    # Open the .docx file
    document = Document(filename)
    fullText = []

    # Iterate through all paragraphs and append their text
    for para in document.paragraphs:
        fullText.append(para.text)

    # Join the list of paragraph strings with newlines
    resume_text = '\n'.join(fullText)

    # if job description is provided, tailor resume statement to it

    if description:
        resume_statement = chat.send_message(
            f"Read the following text denoted by triple backticks and generate a single resume statement that is 70 words or less based on the "
            f"job description denoted by double backticks. ```{resume_text}``` ``{description}``"
        )
        cover_letter = chat.send_message(
            f"Read the following text denoted by triple backticks and generate a cover letter based on the"
            f"job description denoted by double backticks. ```{resume_text}``` ``{description}``."
            f"The company name is {company_name}."
            f"The company address is {company_address}."
            f"The date is {datetime.date.today()}"
            f"Today's date is {datetime.date.today()}"
            f"My name is Hai Bao"
            f"My address is 25 Viking Lane, Unit 1846, Etobicoke, ON, M9B0A1"
            f"My phone number is 647-831-8623"
            f"My email address is hbao12@gmail.com"
        )
        cover_letter_text = cover_letter.text.replace("[Company Name]",company_name).replace("[Company Address]",company_address).replace("*","")

    else:
        resume_statement = chat.send_message(
            f"Read the following text denoted by triple backticks and generate a single resume statement that is 70 words or less.```{resume_text}```")

    # replacement default statement with generated statement
    document.paragraphs[7].text = resume_statement.text.replace("*","")
    document.paragraphs[7].runs[0].font.color.rgb = RGBColor(66, 66, 66)
    document.paragraphs[7].runs[0].font.size = 114300
    document.paragraphs[7].runs[0].font.name = "Merriweather"
    document.paragraphs[7].runs[0].bold = False

    # date string
    date_string = datetime.datetime.now().strftime("%Y_%m_%d")

    # create resume
    document.save(f'resume/resume_{company_code}_{date_string}.docx')

    # create cover letter
    if description:
        cover_letter_docx = Document()
        cover_letter_docx.add_paragraph(cover_letter_text)
        cover_letter_docx.save(f'resume/coverletter_{company_code}_{date_string}.docx')

    # create html version of resume
    resume_file_name = f"resume_{company_code}_{date_string}"
    if not description:
        resume_file_name = f"resume_None_{date_string}"
    html_content = docx_to_html(f"resume/{resume_file_name}.docx").replace("font-size: 9pt","font-size: 14pt").replace("font-size: 11pt", "font-size: 14pt").replace("font-size: 10pt", "font-size:16pt").replace("body { padding: 14pt 36pt 17pt 36pt; }","body { padding: 14pt 36pt 17pt 36pt; margin-left: 0; width: 50% }")
    soup = BeautifulSoup(html_content, "html.parser")
    html_content = soup.prettify()
    if description:
        with open(f"resume/resume_{company_code}_{date_string}.html", "w", encoding="utf-8") as file:
            file.write(html_content)
    else:
        with open(f"resume.html", "w", encoding="utf-8") as file:
            file.write(html_content)

    # create pdf version of resume
    convert_docx_to_pdf(f"resume/resume_{company_code}_{date_string}.docx")

def main(company_name=None, company_address=None, company_code=None, job_description=None):

    file_path = "Resume_ATSF.docx"
    create_resume(file_path, company_name, company_address, company_code, job_description)

if __name__ == "__main__":
    job_description = """
Job Title: Python Developer
    
We are seeking a Python Developer to join a diverse team and help build the next generation of Pricing tools that will redefine the way Aviva commercial underwriters interact with customers, brokers, and vendors.

Aviva Canada recently purchased new Pricing tool, namely Hyperexponential to help build, deploy, and refine pricing models at ease. The platform is native Python and developers can deploy in few clicks raters to production. You will actively develop numerous but uniquely designed HX products for our 20+ commercial lines of businesses within the Global Corporate and Specialty (GCS) space. We aim to onboard all GCS lines of businesses onto the platform within a 2-year timeframe.

You will work closely with other python developers, data scientists, and pricing actuaries to bring to life HX products embedding pricing models to market. You will also work with other stakeholders on the business front such as commercial underwriters to gather requirements for each rater. In addition, you will interact with IT professionals to integrate HX products to other technologies within the pricing ecosystem using APIs and micro services capability.

What you’ll do: 
Develop and maintain HX pricing products and deploy into production in collaboration with other developers, business stakeholders and IT professionals.

Integrate HX products to the wider pricing ecosystem such as Policy Administrative System and Datalake using APIs and micro services.  

Design, build, and maintain efficient, reusable, reliable and secure python codebase using TDD principles.

Mentor development resources on design and code best practices, quality, performance, and security.

Share best practices with team members and internal developer community as well as collaboration and stakeholder management with the broader team.

Work in a team using Scrum, Kanban rituals with a passion for Agile methodology but the flexibility to think outside the scrum.

What you’ll bring:
3+ years of experience in Python Development or software engineering.

Good grasp of software engineering practices such as code-reusability, modularity, etc.

An educational background in computer science or engineering, math, statistics, physics, or a related field. A minimum of a BSc is required, Masters is preferred.

Work collaboratively with other developers in a version control environment using Git practices and repository structure.

Intuitive problem-solving and analytical skills

Strong communication and collaboration skills.

What makes you stand out:
Proficiency with SQL, Datalake, Snowflake and AWS

    """
    company_name = "Aviva Insurance (Canada)"
    company_code = "AVIVA_PD"
    company_address = chat.send_message(
        f"Get the company address for the following company: {company_name}"
    ).text
    main(company_name, company_address, company_code, job_description)
    #main()



